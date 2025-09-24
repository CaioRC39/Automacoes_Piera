# ==============================================================================
# SCRIPT STREAMLIT - GERADOR DE RELATÓRIO DE PROJETOS DE PD&I
# ==============================================================================

# ------------------------------------------------------------------------------
# 1. IMPORTAÇÃO DAS BIBLIOTECAS
# ------------------------------------------------------------------------------
import streamlit as st
import pandas as pd
import docx
import io
import os
import openpyxl
import re
import pypandoc
from pypandoc.pandoc_download import download_pandoc
from openpyxl.styles import Font, PatternFill, Alignment

try:
    pypandoc.get_pandoc_path()
except OSError:
    download_pandoc()

# ------------------------------------------------------------------------------
# 2. FUNÇÕES AUXILIARES
# ------------------------------------------------------------------------------
@st.cache_data
def extract_lp_data_from_docx(doc_content_bytes):
    try:
        doc = docx.Document(io.BytesIO(doc_content_bytes))
        resultados = {}

         # --- 1. EXTRAÇÃO DIRETA DO DOCX (TEXTOS E COORDENADAS) ---
        resultados["Nome do Projeto"] = doc.tables[0].cell(1, 0).text.strip()
        resultados["Descrição do Projeto"] = doc.tables[2].cell(0, 0).text.strip()
        resultados["Justificativa TRL"] = doc.tables[5].cell(0, 0).text.strip()
        resultados["Elemento Inovador"] = doc.tables[9].cell(0, 0).text.strip()
        resultados["Barreiras/Desafios"] = doc.tables[10].cell(0, 0).text.strip()
        resultados["Metodologias"] = doc.tables[11].cell(0, 0).text.strip()
        resultados["Atividades Ano-Base"] = doc.tables[14].cell(0, 0).text.strip()
        resultados["Informações complementares"] = doc.tables[15].cell(0, 0).text.strip()
        resultados["Resultado Econômico"] = doc.tables[16].cell(0, 0).text.strip()
        resultados["Resultado de inovação"] = doc.tables[17].cell(0, 0).text.strip()
        resultados["Justificativa ODS"] = doc.tables[19].cell(0, 0).text.strip()
        resultados["Alinhamento Políticas (Justificativa)"] = doc.tables[20].cell(0, 0).text.strip()

        def find_value(label):
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) > 1 and label in row.cells[0].text: return row.cells[1].text.strip()
            return ""
        trl_inicial_texto = find_value("TRL Inicial:")
        trl_final_texto = find_value("TRL Final:")
        resultados["TRL Inicial"] = re.search(r'\d+', trl_inicial_texto).group(0) if re.search(r'\d+', trl_inicial_texto) else ""
        resultados["TRL Final"] = re.search(r'\d+', trl_final_texto).group(0) if re.search(r'\d+', trl_final_texto) else ""
        resultados["Data de início"] = find_value("Data de início (dia/mês/ano):")
        resultados["Data de término"] = find_value("Data de término (dia/mês/ano):")
        palavras = [row.cells[1].text.strip() for row in doc.tables[8].rows if "Palavra-chave" in row.cells[0].text and len(row.cells) > 1 and row.cells[1].text.strip()]
        resultados["Palavras-chave"] = ", ".join(palavras)

	# --- 2. EXTRAÇÃO DE TODOS OS CHECKBOXES (VIA CONVERSÃO DOCX->TXT) ---
        temp_path = 'temp_doc_for_conversion.docx'
        with open(temp_path, 'wb') as f: f.write(doc_content_bytes)
        plain_text = pypandoc.convert_file(temp_path, 'plain', format='docx', extra_args=['--wrap=none'])

        def get_section_text(full_text, start_keyword, end_keyword):
            try:
                start_index = full_text.lower().find(start_keyword.lower())
                end_index = full_text.lower().find(end_keyword.lower(), start_index)
                if start_index == -1 or end_index == -1: return ""
                return full_text[start_index:end_index]
            except: return ""
        def find_checked_in_section(section_text, is_list=False):
            found = re.findall(r'☒\s*([^☐☒\n\t]+)', section_text)
            cleaned_found = [item.replace('*', '').strip() for item in found]
            if is_list: return cleaned_found
            else: return cleaned_found[0] if cleaned_found else ""
        class_texto = find_checked_in_section(get_section_text(plain_text, "Classificação da pesquisa", "TRL Inicial"))
        if "Pesquisa básica dirigida" in class_texto: resultados["Classificação (PB, PA, DE)"] = "PB"
        elif "Pesquisa aplicada" in class_texto: resultados["Classificação (PB, PA, DE)"] = "PA"
        elif "Desenvolvimento experimental" in class_texto: resultados["Classificação (PB, PA, DE)"] = "DE"
        else: resultados["Classificação (PB, PA, DE)"] = ""
        natureza_texto = find_checked_in_section(get_section_text(plain_text, "Natureza Predominante", "Elemento Tecnologicamente Novo"))
        if "Processos Empresariais" in natureza_texto: resultados["Natureza"] = "Processo"
        elif "Produto - Bens" in natureza_texto: resultados["Natureza"] = "Produto"
        elif "Produto - Serviços" in natureza_texto: resultados["Natureza"] = "Serviço"
        else: resultados["Natureza"] = ""
        resultados["Atividade Contínua"] = find_checked_in_section(get_section_text(plain_text, "A atividade é contínua", "ATIVIDADES DE P,D&I"))
        resultados["Alinhamento Políticas (Sim/Não)"] = find_checked_in_section(get_section_text(plain_text, "políticas públicas nacionais", "Alinhamento do Projeto com Políticas"))
        area_encontrada = find_checked_in_section(get_section_text(plain_text, "Área do projeto", "Palavras-Chave"), is_list=True)
        resultados["Área do projeto"] = ", ".join(area_encontrada)
        ods_encontrados_texto = find_checked_in_section(get_section_text(plain_text, "Objetivos de Desenvolvimento Sustentável", "Justificativa (ODS)"), is_list=True)
        ods_numeros = [re.search(r'\d+', ods).group(0) for ods in ods_encontrados_texto if re.search(r'\d+', ods)]
        resultados["ODS"] = ", ".join(ods_numeros)
        return resultados
    except Exception as e:
        st.error(f"Erro ao extrair dados do Word: {e}")
        return {}

@st.cache_data
def load_sheet_with_dynamic_header(_file_content_bytes, sheet_name, keyword='LINHA DE PESQUISA'):
    try:
        df_no_header = pd.read_excel(io.BytesIO(_file_content_bytes), sheet_name=sheet_name, header=None)
        header_row_index = next((i for i, r in df_no_header.head(20).iterrows() if any(str(c).strip().upper() == keyword.upper() for c in r.values)), -1)
        if header_row_index == -1: raise ValueError(f"Cabeçalho com '{keyword}' não encontrado na aba '{sheet_name}'.")
        df = pd.read_excel(io.BytesIO(_file_content_bytes), sheet_name=sheet_name, header=header_row_index)
        df.columns = [re.sub(r'\s+', ' ', str(col)).strip() for col in df.columns]
        return df
    except Exception as e:
        st.error(f"Erro ao carregar a aba '{sheet_name}': {e}")
        return pd.DataFrame()

def aplicar_formatacao_final(writer):
    workbook = writer.book
    header_font = Font(bold=True, color="000000")
    header_fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")
    formato_duas_casas = '0.00'
    for sheet_name in ['RH', 'ST', 'LP']:
        if sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            for cell in worksheet[1]:
                cell.font = header_font
                cell.fill = header_fill
                cell.alignment = header_alignment
            for col_idx, column_cell in enumerate(worksheet[1], 1):
                if column_cell.value in ["VALOR TOTAL", "HORAS"]:
                    for cell in worksheet.iter_cols(min_col=col_idx, max_col=col_idx, min_row=2):
                        cell[0].number_format = formato_duas_casas

# ------------------------------------------------------------------------------
# 3. INTERFACE DO STREAMLIT
# ------------------------------------------------------------------------------

st.set_page_config(layout="wide", page_title="Extrator de LP&RH&ST")
st.title("📊 Extrator de Informações da Valoração")
import streamlit as st

# Título principal
st.markdown("## **Extração de dados de Linha de Pesquisa, Recursos Humanos e Serviços de Terceiros e Viagens**")

# Descrição Geral
st.markdown("## **Descrição Geral**")
st.markdown("""
Este script é uma ferramenta de automação desenvolvida com objetivo de ler e processar dados de múltiplas fontes para gerar um arquivo consolidado em Excel.

A automação elimina o trabalho manual de copiar, colar, filtrar, agrupar e somar dados, garantindo consistência, rapidez e precisão na extração dos dados das abas **LP (Linha de Pesquisa)**, **RH (Recursos Humanos)** e **ST (Serviços de Terceiros)**.
""")

# Pré-requisitos
st.markdown("## **Pré-requisitos**")
st.markdown("""
Para que o script funcione corretamente, você precisará de dois arquivos de entrada:

1. **Planilha de Valoração (`.xlsx`):**
    * Este é o arquivo principal que contém os dados brutos de RH e ST.
    * Deve conter uma aba cujo nome começa com **`Timesheet_`**. Esta aba precisa ter as colunas `LINHA DE PESQUISA`, `PROJETO`, `NOME DO COLABORADOR`, `C.P.F.`, `CARGO`, `HORAS APROPRIADAS A HORAS ÚTEIS` e `LEI DO BEM` completas, principalmente a `LINHA DE PESQUISA`.
    * Deve conter uma aba chamada **`Serviços de Terceiros e Viagens`**. Esta aba precisa ter as colunas `LINHA DE PESQUISA`, `PROJETO`, `RAZÃO SOCIAL PRESTADOR`, `CNPJ PRESTADOR`, `PORTE DA EMPRESA`, `R$ FINAL` e `DESPESA VÁLIDA PARA O PIT?`.

2. **TAs (`.docx`):**
    * São os arquivos que contêm as informações descritivas de cada Linha de Pesquisa.
    * **Ponto Crítico:** O nome de cada arquivo deve corresponder **exatamente** ao nome utilizado na coluna `LINHA DE PESQUISA` da Planilha de Valoração.
""")

# Instruções de Uso
st.markdown("## **Instruções de Uso**")
st.markdown("""
Siga este passo a passo para gerar a extração:

1. **Inserir Nome da Empresa:**
    * O primeiro campo de texto solicita o **Nome da Empresa**. Digite o nome que você deseja que apareça no arquivo final e pressione `Enter`.

3. **Upload dos Arquivos:**
    * O script solicitará os arquivos em dois quadros:
        * **Primeiro, a Planilha de Valoração:** Clique em "Browse files" e selecione a planilha `.xlsx` de Valoração.
        * **Depois, os TAs:** No quadro abaixo, navegue até a pasta onde estão os TAs em `.docx` e selecione **todos** os que deseja processar.
            > **Dica:** Para selecionar múltiplos arquivos, segure a tecla `Ctrl` (no Windows) ou `Cmd` (no Mac) enquanto clica em cada arquivo.

4. **Aguardar o Processamento:**
    * Após o upload, o script começará a processar os dados automaticamente. Você verá mensagens de status na tela informando o progresso para cada Linha de Pesquisa e cada aba.

5. **Download do Relatório:**
    * Ao final do processo, uma mensagem de "Processo Concluído!" e o botão de download do seu novo arquivo Excel aparecerá. Clique nele para baixar a extração.
    * O arquivo final terá o nome no formato: `NOME_DA_EMPRESA_LP&RH&ST.xlsx`.
""")

# O Arquivo de Saída
st.markdown("## **O Arquivo de Saída**")
st.markdown("""
O relatório em Excel gerado conterá três abas:

* **`LP`:** Similar ao NewPiit, uma linha para cada arquivo Word processado, contendo todos os detalhes extraídos do documento (Descrição, TRLs, Natureza, etc.).
* **`RH`:** Um resumo de todos os colaboradores que trabalharam em cada Linha de Pesquisa, com as horas e valores somados por CPF. Os dados são pré-filtrados para remover estagiários e valores zerados.
* **`ST`:** Um resumo de todos os prestadores de serviço para cada Linha de Pesquisa, com os valores de "R$ FINAL" somados por CNPJ. Os dados são pré-filtrados para incluir apenas despesas válidas.
""")


with st.sidebar:
    st.header("⚙️ Configurações")
    nome_empresa_input = st.text_input("1. Nome da Empresa para o arquivo final:", placeholder="Ex: Minha Empresa")
    uploaded_valoracao = st.file_uploader("2. Faça o upload da Planilha de Valoração (.xlsx)", type=['xlsx'])
    uploaded_words = st.file_uploader("3. Faça o upload dos Documentos Word (TAs) (.docx)", type=['docx'], accept_multiple_files=True)
    processar_button = st.button("Gerar Relatório", type="primary", use_container_width=True)

if processar_button:
    if not nome_empresa_input or not uploaded_valoracao or not uploaded_words:
        st.warning("⚠️ Por favor, preencha o nome da empresa e faça o upload de todos os arquivos necessários.")
    else:
        with st.spinner("Processando... Isso pode levar alguns minutos, dependendo do número de arquivos."):
            try:
                nome_empresa_safe = nome_empresa_input.replace(' ', '_')
                valoracao_file_content = uploaded_valoracao.getvalue()

                st.info("1/3 - Processando dados das Linhas de Pesquisa (Word)...")
                novas_linhas_lp = []
                for doc_file in uploaded_words:
                    linha_pesquisa_nome = re.sub(r'\s*\(\d+\)$', '', os.path.splitext(doc_file.name)[0]).strip()
                    lp_data = extract_lp_data_from_docx(doc_file.getvalue())
                    if lp_data:
                        lp_data['Linha de Pesquisa'] = linha_pesquisa_nome
                        novas_linhas_lp.append(lp_data)
                df_lp_final = pd.DataFrame(novas_linhas_lp)
                if not df_lp_final.empty:
                    colunas_lp = [
            'Linha de Pesquisa',
            'Nome do Projeto',
            'Descrição do Projeto',
            'Classificação (PB, PA, DE)',
            'Área do projeto',
            'Palavras-chave',
            'Natureza',
            'Elemento Inovador',
            'Barreiras/Desafios',
            'Metodologias',
            'Atividade Contínua',
            'Data de início',
            'Data de término',
            'Atividades Ano-Base',
            'Informações complementares',
            'Resultado Econômico',
            'Resultado de inovação',
            'TRL Inicial',
            'TRL Final',
            'Justificativa TRL',
            'ODS',
            'Justificativa ODS',
            'Alinhamento Políticas (Sim/Não)',
            'Alinhamento Políticas (Justificativa)'
        ]
                    df_lp_final = df_lp_final.reindex(columns=colunas_lp).fillna('')

                st.info("2/3 - Processando dados de RH (Valoração)...")
                df_rh_final = pd.DataFrame()
                timesheet_sheet_name = [s for s in pd.ExcelFile(io.BytesIO(valoracao_file_content)).sheet_names if s.startswith('Timesheet_')][0]
                df_rh_raw = load_sheet_with_dynamic_header(valoracao_file_content, timesheet_sheet_name, keyword='LINHA DE PESQUISA')
                if not df_rh_raw.empty:
                    lei_do_bem_col_name = next((col for col in df_rh_raw.columns if "LEI DO BEM" in str(col).upper() and "?" not in str(col)), None)
                    if lei_do_bem_col_name:
                        df_rh_raw[lei_do_bem_col_name] = pd.to_numeric(df_rh_raw[lei_do_bem_col_name], errors='coerce').fillna(0)
                        df_rh_filtrado = df_rh_raw[df_rh_raw[lei_do_bem_col_name] > 0].copy()
                        if not df_rh_filtrado.empty:
                            aggregations_rh = {'PROJETO': 'first', 'NOME DO COLABORADOR': 'first', 'CARGO': 'first', 'HORAS APROPRIADAS A HORAS ÚTEIS': 'sum', lei_do_bem_col_name: 'sum'}
                            df_rh_grouped = df_rh_filtrado.groupby(['LINHA DE PESQUISA', 'C.P.F.']).agg(aggregations_rh).reset_index()
                            df_rh_grouped['DESCRIÇÃO DA ATIVIDADE'] = ''
                            df_rh_grouped['HORAS APROPRIADAS A HORAS ÚTEIS'] = df_rh_grouped['HORAS APROPRIADAS A HORAS ÚTEIS'].round(2)
                            df_rh_grouped[lei_do_bem_col_name] = df_rh_grouped[lei_do_bem_col_name].round(2)
                            df_rh_final = df_rh_grouped.rename(columns={'LINHA DE PESQUISA': 'LP', 'NOME DO COLABORADOR': 'COLABORADOR', 'C.P.F.': 'CPF', 'HORAS APROPRIADAS A HORAS ÚTEIS': 'HORAS', lei_do_bem_col_name: 'VALOR TOTAL'})
                            df_rh_final = df_rh_final[['LP', 'PROJETO', 'COLABORADOR', 'CPF', 'CARGO', 'HORAS', 'VALOR TOTAL', 'DESCRIÇÃO DA ATIVIDADE']]
                
                st.info("3/3 - Processando dados de ST (Valoração)...")
                df_st_final = pd.DataFrame()
                df_st_raw = load_sheet_with_dynamic_header(valoracao_file_content, 'Serviços de Terceiros e Viagens', keyword='LINHA DE PESQUISA')
                if not df_st_raw.empty:
                    df_st_filtrado = df_st_raw[df_st_raw['DESPESA VÁLIDA PARA O PIT?'] == 'Sim'].copy()
                    if not df_st_filtrado.empty:
                        aggregations_st = {'PROJETO': 'first', 'RAZÃO SOCIAL PRESTADOR': 'first', 'R$ FINAL': 'sum'}
                        df_st_grouped = df_st_filtrado.groupby(['LINHA DE PESQUISA', 'CNPJ PRESTADOR']).agg(aggregations_st).reset_index()
                        df_st_grouped['R$ FINAL'] = df_st_grouped['R$ FINAL'].round(2)
                        df_st_grouped['DESCRIÇÃO DA ATIVIDADE'] = ''
                        df_st_final = df_st_grouped.rename(columns={'LINHA DE PESQUISA': 'LP', 'R$ FINAL': 'VALOR TOTAL'})
                        df_st_final = df_st_final[['LP', 'PROJETO', 'RAZÃO SOCIAL PRESTADOR', 'CNPJ PRESTADOR', 'VALOR TOTAL', 'DESCRIÇÃO DA ATIVIDADE']]
                
                st.info("Gerando arquivo Excel final...")
                output_stream = io.BytesIO()
                with pd.ExcelWriter(output_stream, engine='openpyxl') as writer:
                    df_lp_final.to_excel(writer, sheet_name='LP', index=False)
                    df_rh_final.to_excel(writer, sheet_name='RH', index=False)
                    df_st_final.to_excel(writer, sheet_name='ST', index=False)
                    aplicar_formatacao_final(writer)
                
                st.success("🎉 Relatório gerado com sucesso!")
                output_filename = f"{nome_empresa_safe}_LP&RH&ST.xlsx"
                st.download_button(
                    label="📥 Baixar Relatório (.xlsx)",
                    data=output_stream,
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

            except Exception as e:
                st.error(f"Ocorreu um erro durante o processamento: {e}")
