# ==============================================================================
# SCRIPT STREAMLIT - PREENCHEDOR AUTOMÁTICO DE PLANILHA
# ==============================================================================

# ------------------------------------------------------------------------------
# 1. IMPORTAÇÃO DAS BIBLIOTECAS
# ------------------------------------------------------------------------------
import streamlit as st
import pandas as pd
import docx
import io
import os
import openpyxl
import re
import pypandoc
from openpyxl.styles import Font, PatternFill, Alignment
import math

try:
    pypandoc.get_pandoc_path()
except OSError:
    download_pandoc()

# ------------------------------------------------------------------------------
# 2. FUNÇÕES AUXILIARES
# ------------------------------------------------------------------------------
@st.cache_data
def extract_geral_data(doc_content_bytes):
    try:
        doc = docx.Document(io.BytesIO(doc_content_bytes))
        resultados = {}
        resultados["Nome da atividade de PD&I (Nome do projeto igual no GERAL)"] = doc.tables[0].cell(1, 0).text.strip()
        resultados["Descrição do Projeto:"] = doc.tables[2].cell(0, 0).text.strip()
        resultados["Justificativa TRL"] = doc.tables[5].cell(0, 0).text.strip()
        resultados["Destaque o elemento tecnologicamente novo ou inovador da atividade: \xa0"] = doc.tables[9].cell(0, 0).text.strip()
        resultados["Qual a barreira ou desafio tecnológico superável: \xa0"] = doc.tables[10].cell(0, 0).text.strip()
        resultados["Qual a metodologia / métodos utilizados: \xa0"] = doc.tables[11].cell(0, 0).text.strip()
        resultados["Caso a atividade/projeto seja continuada, informar Atividade de PD&I desenvolvida no ano-base"] = doc.tables[14].cell(0, 0).text.strip()
        resultados["Descrição Complementar: "] = doc.tables[15].cell(0, 0).text.strip()
        resultados["Resultado Econômico:"] = doc.tables[16].cell(0, 0).text.strip()
        resultados["Resultado de Inovação:"] = doc.tables[17].cell(0, 0).text.strip()
        resultados["Justificativa ODS"] = doc.tables[19].cell(0, 0).text.strip()
        resultados["Alinhamento do Projeto com Políticas, Programas e Estratégias Governamentais"] = doc.tables[20].cell(0, 0).text.strip()
        def find_value(label):
            for table in doc.tables:
                for row in table.rows:
                    if len(row.cells) > 1 and label in row.cells[0].text: return row.cells[1].text.strip()
            return ""
        trl_inicial_texto = find_value("TRL Inicial:")
        trl_final_texto = find_value("TRL Final:")
        resultados["TRL Inicial"] = re.search(r'\d+', trl_inicial_texto).group(0) if re.search(r'\d+', trl_inicial_texto) else ""
        resultados["TRL Final"] = re.search(r'\d+', trl_final_texto).group(0) if re.search(r'\d+', trl_final_texto) else ""
        resultados["Data de início: (formato dd/mm/aaaa)"] = find_value("Data de início (dia/mês/ano):")
        resultados["Previsão de término: (formato dd/mm/aaaa)"] = find_value("Data de término (dia/mês/ano):")
        palavras = [row.cells[1].text.strip() for row in doc.tables[8].rows if "Palavra-chave" in row.cells[0].text and len(row.cells) > 1 and row.cells[1].text.strip()]
        resultados["Palavras-Chave (Separadas por vírgula):"] = ", ".join(palavras)
        def find_checked_para(options):
            for p in doc.paragraphs:
                if '<w14:checked w14:val="1"/>' in p._element.xml:
                    for opt in options:
                        if opt in p.text: return opt
            return ""
        classificacao_texto = find_checked_para(["Pesquisa básica dirigida", "Pesquisa aplicada", "Desenvolvimento experimental"])
        if "Pesquisa básica dirigida" in classificacao_texto: resultados["PB, PA ou DE:"] = "PB"
        elif "Pesquisa aplicada" in classificacao_texto: resultados["PB, PA ou DE:"] = "PA"
        elif "Desenvolvimento experimental" in classificacao_texto: resultados["PB, PA ou DE:"] = "DE"
        else: resultados["PB, PA ou DE:"] = ""
        natureza_texto = find_checked_para(["Processos Empresariais", "Produto - Bens", "Produto - Serviços"])
        if "Processos Empresariais" in natureza_texto: resultados["Natureza (Produto, Processo ou Serviço):"] = "Processo"
        elif "Produto - Bens" in natureza_texto: resultados["Natureza (Produto, Processo ou Serviço):"] = "Produto"
        elif "Produto - Serviços" in natureza_texto: resultados["Natureza (Produto, Processo ou Serviço):"] = "Serviço"
        else: resultados["Natureza (Produto, Processo ou Serviço):"] = ""
        resultados["Os projetos de PD&I da empresa se alinham com as políticas públicas nacionais? (Sim ou Não)"] = find_checked_para(["Sim", "Não"])
        resultados["A atividade é contínua (ciclo de vida maior que 1 ano)?\xa0 (Sim ou Não)"] = find_checked_para(["Sim", "Não"])
        temp_path = 'temp_doc_for_conversion.docx'
        with open(temp_path, 'wb') as f: f.write(doc_content_bytes)
        plain_text = pypandoc.convert_file(temp_path, 'plain', format='docx', extra_args=['--wrap=none'])
        def get_section_text(full_text, start_keyword, end_keyword):
            try:
                start_index = full_text.lower().find(start_keyword.lower())
                end_index = full_text.lower().find(end_keyword.lower(), start_index)
                if start_index == -1 or end_index == -1: return ""
                return full_text[start_index:end_index]
            except: return ""
        area_texto = get_section_text(plain_text, "Área do projeto", "Palavras-Chave")
        ods_texto = get_section_text(plain_text, "Objetivos de Desenvolvimento Sustentável", "Justificativa (ODS)")
        area_encontrada = re.findall(r'☒\s*([A-ZÀ-Ú][^☐☒\n]+)', area_texto)
        ods_encontrados_texto = re.findall(r'☒\s*(\d+\.\s*[^☐☒\n]+)', ods_texto)
        ods_numeros = [re.search(r'\d+', ods).group(0) for ods in ods_encontrados_texto if re.search(r'\d+', ods)]
        resultados["Área do Projeto:"] = ", ".join([area.strip() for area in area_encontrada])
        resultados["ODS"] = ", ".join(ods_numeros)
        if resultados.get("A atividade é contínua (ciclo de vida maior que 1 ano)?\xa0 (Sim ou Não)") == "Não":
            resultados["Data de início: (formato dd/mm/aaaa)"] = ""
            resultados["Previsão de término: (formato dd/mm/aaaa)"] = ""
            resultados["Caso a atividade/projeto seja continuada, informar Atividade de PD&I desenvolvida no ano-base"] = ""
        if resultados.get("Os projetos de PD&I da empresa se alinham com as políticas públicas nacionais? (Sim ou Não)") == "Não":
            resultados["Alinhamento do Projeto com Políticas, Programas e Estratégias Governamentais"] = ""
        return resultados
    except Exception as e:
        st.error(f"Erro ao extrair dados do Word: {e}")
        return {}

@st.cache_data
def load_sheet_with_dynamic_header(_file_content_bytes, sheet_name, keyword='LINHA DE PESQUISA'):
    try:
        df_no_header = pd.read_excel(io.BytesIO(_file_content_bytes), sheet_name=sheet_name, header=None)
        header_row_index = next((i for i, r in df_no_header.head(20).iterrows() if any(str(c).strip().upper() == keyword.upper() for c in r.values)), -1)
        if header_row_index == -1: raise ValueError(f"Cabeçalho com '{keyword}' não encontrado na aba '{sheet_name}'.")
        df = pd.read_excel(io.BytesIO(_file_content_bytes), sheet_name=sheet_name, header=header_row_index)
        df.columns = [re.sub(r'\s+', ' ', str(col)).strip() for col in df.columns]
        return df
    except Exception as e:
        st.error(f"Erro ao carregar a aba '{sheet_name}': {e}")
        return pd.DataFrame()

# ------------------------------------------------------------------------------
# 3. INTERFACE DO STREAMLIT
# ------------------------------------------------------------------------------

st.set_page_config(layout="wide", page_title="Preenchedor Automático")
st.title("📄 Preenchedor Automático de Planilha")
st.markdown("Esta ferramenta preenche o NewPiit a partir da planilha de Valoração e dos TAs.")

with st.sidebar:
    st.header("⚙️ Configurações")
    nome_empresa_input = st.text_input("1. Nome da Empresa para o arquivo final:", placeholder="Ex: Minha Empresa")
    uploaded_base = st.file_uploader("2. Faça o upload do NewPiit (.xlsx)", type=['xlsx'])
    uploaded_valoracao = st.file_uploader("3. Faça o upload da Planilha de Valoração (.xlsx)", type=['xlsx'])
    uploaded_words = st.file_uploader("4. Faça o upload dos TAs (.docx)", type=['docx'], accept_multiple_files=True)
    processar_button = st.button("Preencher Planilha", type="primary", use_container_width=True)

if processar_button:
    if not all([nome_empresa_input, uploaded_base, uploaded_valoracao, uploaded_words]):
        st.warning("⚠️ Por favor, preencha o nome da empresa e faça o upload de todos os arquivos necessários.")
    else:
        with st.spinner("Processando... Isso pode levar alguns minutos."):
            try:
                # Lógica principal do script
                nome_empresa_safe = nome_empresa_input.replace(' ', '_')
                base_file_content = uploaded_base.getvalue()
                valoracao_file_content = uploaded_valoracao.getvalue()
                base_filename_cleaned = re.sub(r'\s*\(\d+\)$', '', os.path.splitext(uploaded_base.name)[0]).strip()

                st.info("Carregando planilha de Valoração...")
                df_disp = load_sheet_with_dynamic_header(valoracao_file_content, 'Serviços de Terceiros e Viagens', keyword='LINHA DE PESQUISA')
                timesheet_name = [s for s in pd.ExcelFile(io.BytesIO(valoracao_file_content)).sheet_names if s.startswith('Timesheet_')][0]
                df_rh = load_sheet_with_dynamic_header(valoracao_file_content, timesheet_name, keyword='LINHA DE PESQUISA')
                df_disp['LINHA DE PESQUISA'] = df_disp['LINHA DE PESQUISA'].astype(str).str.strip()
                df_rh['LINHA DE PESQUISA'] = df_rh['LINHA DE PESQUISA'].astype(str).str.strip()

                # Bloco de Preparação para Validação
                mapa_lp_para_projetos, gabarito_totais = {}, {}
                if not df_rh.empty and 'LINHA DE PESQUISA' in df_rh.columns and 'PROJETO' in df_rh.columns:
                    df_rh_clean = df_rh.dropna(subset=['LINHA DE PESQUISA', 'PROJETO'])
                    mapeamento = df_rh_clean[['LINHA DE PESQUISA', 'PROJETO']].drop_duplicates()
                    for _, row in mapeamento.iterrows():
                        lp, proj = str(row['LINHA DE PESQUISA']).strip(), str(row['PROJETO']).strip()
                        if lp not in mapa_lp_para_projetos: mapa_lp_para_projetos[lp] = []
                        if proj not in mapa_lp_para_projetos[lp]: mapa_lp_para_projetos[lp].append(proj)
                try:
                    resumo_sheet_name = [s for s in pd.ExcelFile(io.BytesIO(valoracao_file_content)).sheet_names if s.startswith('Resumo')][0]
                    df_resumo = pd.read_excel(io.BytesIO(valoracao_file_content), sheet_name=resumo_sheet_name, header=None)
                    for _, row in df_resumo.iterrows():
                        try:
                            nome_projeto = str(row.iloc[2]).strip()
                            if nome_projeto and nome_projeto.lower() not in ["total", "projeto"]:
                                total_rh, total_st = float(row.iloc[4]), float(row.iloc[5])
                                if nome_projeto not in gabarito_totais: gabarito_totais[nome_projeto] = {'RH': 0, 'ST': 0}
                                gabarito_totais[nome_projeto]['RH'] += total_rh
                                gabarito_totais[nome_projeto]['ST'] += total_st
                        except (ValueError, IndexError): continue
                except Exception as e: st.warning(f"Não foi possível ler totais da aba Resumo. Erro: {e}")

                novas_linhas_geral, novas_linhas_disp_st, novas_linhas_rh = [], [], []
                id_disp, id_rh = 1, 1

                progress_bar = st.progress(0, text="Processando arquivos Word...")
                for idx, doc_file in enumerate(uploaded_words):
                    nome_busca_projeto = re.sub(r'\s*\(\d+\)$', '', os.path.splitext(doc_file.name)[0]).strip()
                    st.info(f"Processando Linha de Pesquisa: '{nome_busca_projeto}'")
                    geral_data_extraida = extract_geral_data(doc_file.getvalue())
                    if geral_data_extraida:
                        nome_final_projeto = geral_data_extraida.get('Nome da atividade de PD&I (Nome do projeto igual no GERAL)', nome_busca_projeto)
                        geral_data_extraida['#'] = idx + 1
                        geral_data_extraida['Nome da atividade de PD&I (Nome do projeto igual no GERAL)'] = nome_final_projeto
                        novas_linhas_geral.append(geral_data_extraida)
                        
                        # Processamento ST
                        df_f_st = df_disp[(df_disp['LINHA DE PESQUISA'] == nome_busca_projeto) & (df_disp['DESPESA VÁLIDA PARA O PIT?'] == 'Sim')]
                        if not df_f_st.empty:
                            for cnpj, grupo in df_f_st.groupby('CNPJ PRESTADOR'):
                                novas_linhas_disp_st.append({'#': id_disp, 'Nome da atividade de PD&I (Nome do projeto igual no GERAL)': nome_final_projeto,'TIPO': str(grupo.iloc[0]['PORTE DA EMPRESA']).title(), 'Situação (Contratado, Em Execução, Terminado)': 'Terminado','Prestador de Serviço': grupo.iloc[0]['RAZÃO SOCIAL PRESTADOR'], 'CNPJ/CPF': cnpj,'Caracterizar o Serviço Realizado': 'Serviço de apoio técnico para desenvolvimento do projeto','Valor Total': round(grupo['R$ FINAL'].sum(), 2)}); id_disp += 1
                        
                        # Processamento RH
                        lei_do_bem_col = next((c for c in df_rh.columns if "LEI DO BEM" in c.upper() and "?" not in c), None)
                        if lei_do_bem_col:
                            df_rh[lei_do_bem_col] = pd.to_numeric(df_rh[lei_do_bem_col], errors='coerce').fillna(0)
                            df_f_rh = df_rh[(df_rh['LINHA DE PESQUISA'] == nome_busca_projeto) & (df_rh[lei_do_bem_col] != 0) & (~df_rh['CARGO'].str.contains('Estagiario', case=False, na=False))].copy()
                            if not df_f_rh.empty:
                                def categorizar_escolaridade(texto):
                                    texto_limpo_title, texto_limpo_lower = str(texto).strip().title(), str(texto).lower().strip()
                                    lista_validos = ["Doutor", "Mestre", "Pós-Graduado", "Graduado", "Tecnólogo", "Técnico De Nível Médio", "Apoio Técnico"]
                                    if texto_limpo_title in lista_validos: return texto_limpo_title
                                    if any(s in texto_limpo_lower for s in ['especialização', 'pós-graduado']): return 'Pós-graduado'
                                    if any(s in texto_limpo_lower for s in ['superior completa', 'superior completo']): return 'Graduado'
                                    if any(s in texto_limpo_lower for s in ['superior incompleta', 'superior incompleto', 'médio completo']): return 'Apoio Técnico'
                                    return "Apoio Técnico"
                                df_f_rh['TITULAÇÃO_CONVERTIDA'] = df_f_rh['ESCOLARIDADE'].apply(categorizar_escolaridade)
                                for cpf, grupo in df_f_rh.groupby('C.P.F.'):
                                    novas_linhas_rh.append({'#': id_rh, 'Nome da atividade de PD&I (Nome do projeto igual no GERAL)': nome_final_projeto, 'CPF': cpf,'NOME': grupo.iloc[0]['NOME DO COLABORADOR'], 'TITULAÇÃO': grupo.iloc[0]['TITULAÇÃO_CONVERTIDA'],'Total Horas (Anual)': round(grupo['HORAS APROPRIADAS A HORAS ÚTEIS'].sum(), 2),'Valor (R$)': round(grupo[lei_do_bem_col].sum(), 2)}); id_rh += 1
                    progress_bar.progress((idx + 1) / len(uploaded_words), text=f"Processando {doc_file.name}...")
                
                # Bloco de Validação Final
                st.info("Validando totais calculados...")
                validation_messages = []
                for lp_nome_arquivo in uploaded_words:
                    lp_limpo = re.sub(r'\s*\(\d+\)$', '', os.path.splitext(lp_nome_arquivo.name)[0]).strip()
                    projetos_na_lp = mapa_lp_para_projetos.get(lp_limpo, [])
                    if not projetos_na_lp:
                        validation_messages.append(f"**{lp_limpo}:** AVISO - Relação entre Linha de Pesquisa e Projetos não encontrada.")
                        continue
                    soma_esperada_st = sum(gabarito_totais.get(proj, {}).get('ST', 0) for proj in projetos_na_lp)
                    soma_esperada_rh = sum(gabarito_totais.get(proj, {}).get('RH', 0) for proj in projetos_na_lp)
                    nome_final_projeto_atual = next((l['Nome da atividade de PD&I (Nome do projeto igual no GERAL)'] for l in novas_linhas_geral if l['#'] == uploaded_words.index(lp_nome_arquivo) + 1), lp_limpo)
                    soma_calculada_st = sum(l['Valor Total'] for l in novas_linhas_disp_st if l['Nome da atividade de PD&I (Nome do projeto igual no GERAL)'] == nome_final_projeto_atual)
                    soma_calculada_rh = sum(l['Valor (R$)'] for l in novas_linhas_rh if l['Nome da atividade de PD&I (Nome do projeto igual no GERAL)'] == nome_final_projeto_atual)
                    
                    msg_st = f"ST: Calculado ({soma_calculada_st:.2f}) vs Esperado ({soma_esperada_st:.2f})"
                    msg_rh = f"RH: Calculado ({soma_calculada_rh:.2f}) vs Esperado ({soma_esperada_rh:.2f})"
                    status_st = "✅" if math.isclose(soma_calculada_st, soma_esperada_st, rel_tol=0.01) else "⚠️ ALERTA"
                    status_rh = "✅" if math.isclose(soma_calculada_rh, soma_esperada_rh, rel_tol=0.01) else "⚠️ ALERTA"
                    validation_messages.append(f"**{lp_limpo}:** {status_st} {msg_st} | {status_rh} {msg_rh}")

                with st.expander("Ver Resultados da Validação de Totais"):
                    for msg in validation_messages:
                        st.markdown(msg)

                # Geração do Arquivo Excel Final
                output_stream = io.BytesIO()
                wb = openpyxl.load_workbook(io.BytesIO(base_file_content))
                def clear_and_write(sheet_name, data, header_row=10, start_row=11):
                    if data and sheet_name in wb.sheetnames:
                        ws = wb[sheet_name]
                        template_styles = [cell._style for cell in (ws[start_row] if ws.max_row >= start_row else ws[header_row])]
                        if ws.max_row >= start_row:
                            for row in ws.iter_rows(min_row=start_row, max_row=ws.max_row):
                                for cell in row: cell.value = None
                        df = pd.DataFrame(data)
                        header = [str(cell.value).strip() for cell in ws[header_row]]
                        df.columns = [str(col).strip() for col in df.columns]
                        df_ordered = df.reindex(columns=header).fillna('')
                        for r_idx, row_data in enumerate(df_ordered.itertuples(index=False), start_row):
                            for c_idx, value in enumerate(row_data, 1):
                                cell = ws.cell(row=r_idx, column=c_idx)
                                if c_idx - 1 < len(template_styles): cell._style = template_styles[c_idx - 1]
                                cell.value = value
                
                mapa_final = {'Nome da atividade de PD&I (Nome do projeto igual no GERAL)': 'Nome da atividade de PD&I: \xa0','Descrição do Projeto:': 'Descrição do Projeto:','PB, PA ou DE:': 'PB, PA ou DE:','Área do Projeto:': 'Área do Projeto:','Palavras-Chave (Separadas por vírgula):': 'Palavras-Chave (Separadas por vírgula):','Natureza (Produto, Processo ou Serviço):': 'Natureza (Produto, Processo ou Serviço):','Destaque o elemento tecnologicamente novo ou inovador da atividade: \xa0': 'Destaque o elemento tecnologicamente novo ou inovador da atividade: \xa0','Qual a barreira ou desafio tecnológico superável: \xa0': 'Qual a barreira ou desafio tecnológico superável: \xa0','Qual a metodologia / métodos utilizados: \xa0': 'Qual a metodologia / métodos utilizados: \xa0','A atividade é contínua (ciclo de vida maior que 1 ano)?\xa0 (Sim ou Não)': 'A atividade é contínua (ciclo de vida maior que 1 ano)?\xa0 (Sim ou Não)','Data de início: (formato dd/mm/aaaa)': 'Data de início: (formato dd/mm/aaaa)','Previsão de término: (formato dd/mm/aaaa)': 'Previsão de término: (formato dd/mm/aaaa)','Caso a atividade/projeto seja continuada, informar Atividade de PD&I desenvolvida no ano-base': 'Caso a atividade/projeto seja continuada, informar Atividade de PD&I desenvolvida no ano-base','Descrição Complementar: ': 'Descrição Complementar: ','Resultado Econômico:': 'Resultado Econômico:','Resultado de Inovação:': 'Resultado de Inovação:','TRL Inicial': 'TRL Inicial', 'TRL Final': 'TRL Final','Justificativa TRL': 'Justificativa TRL', 'ODS': 'ODS', 'Justificativa ODS': 'Justificativa ODS','Os projetos de PD&I da empresa se alinham com as políticas públicas nacionais? (Sim ou Não)': 'Os projetos de PD&I da empresa se alinham com as políticas públicas nacionais? (Sim ou Não)','Alinhamento do Projeto com Políticas, Programas e Estratégias Governamentais': 'Alinhamento do Projeto com Políticas, Programas e Estratégias Governamentais'}
                df_geral_final = pd.DataFrame(novas_linhas_geral).rename(columns=mapa_final)

                clear_and_write('GERAL', df_geral_final.to_dict('records'))
                clear_and_write('DISPÊNDIOS ST', novas_linhas_disp_st)
                clear_and_write('RH', novas_linhas_rh)
                
                output_filename = f"{nome_empresa_safe}_{base_filename_cleaned}_PREENCHIDO.xlsx"
                wb.save(output_stream)
                
                st.success("🎉 NewPiit preenchido com sucesso!")
                st.download_button(
                    label="📥 Baixar NewPiit Preenchida (.xlsx)",
                    data=output_stream.getvalue(),
                    file_name=output_filename,
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )

            except Exception as e:
                st.error(f"Ocorreu um erro durante o processamento: {e}")
